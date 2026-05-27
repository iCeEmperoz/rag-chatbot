# 01_indexing.py
# Đọc PDF, DOCX, CSV → Chunking → Embed → Lưu vào ChromaDB (Fallback BM25 local if 429)

import os
import sys
import shutil
import json
import warnings
from pathlib import Path
from dotenv import load_dotenv

warnings.filterwarnings("ignore", category=DeprecationWarning)

from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_community.retrievers import BM25Retriever

# Thư viện bổ sung cho Word và CSV
import docx
import pandas as pd

# ── Cấu hình ──────────────────────────────────────────────────
load_dotenv()  # Đọc GOOGLE_API_KEY từ file .env

DATA_DIR    = Path("data")
CHROMA_PATH = Path("chroma_db")
CHUNK_SIZE  = 800
CHUNK_OVERLAP = 100

# Đảm bảo thư mục dữ liệu tồn tại
DATA_DIR.mkdir(exist_ok=True)

# ── Step 1: Load và xử lý đa định dạng (PDF, DOCX, CSV) ────────
def load_all_documents(data_dir: Path) -> list:
    all_chunks = []
    
    # 1. Xử lý PDF
    pdf_files = list(data_dir.glob("*.pdf"))
    if pdf_files:
        print("\n--- Đang xử lý các file PDF ---")
        for pdf_path in pdf_files:
            print(f"  Đang đọc PDF: {pdf_path.name}")
            loader = PyPDFLoader(str(pdf_path))
            pages = loader.load()
            
            # Cắt nhỏ văn bản PDF
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=CHUNK_SIZE,
                chunk_overlap=CHUNK_OVERLAP,
                separators=["\n\n", "\n", ". ", " ", ""],
            )
            pdf_chunks = splitter.split_documents(pages)
            
            # Chuẩn hóa metadata
            for chunk in pdf_chunks:
                page_num = chunk.metadata.get("page", 0) + 1  # PyPDFLoader index từ 0
                chunk.metadata = {
                    "source": pdf_path.name,
                    "page": page_num,
                    "type": "PDF"
                }
            all_chunks.extend(pdf_chunks)
            print(f"  → Đã trích xuất {len(pdf_chunks)} chunks từ PDF.")
    
    # 2. Xử lý DOCX
    docx_files = list(data_dir.glob("*.docx"))
    if docx_files:
        print("\n--- Đang xử lý các file Word (DOCX) ---")
        for docx_path in docx_files:
            print(f"  Đang đọc Word: {docx_path.name}")
            doc = docx.Document(str(docx_path))
            # Ghép các đoạn văn không trống lại với nhau
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)
            
            # Cắt nhỏ văn bản DOCX
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=CHUNK_SIZE,
                chunk_overlap=CHUNK_OVERLAP,
                separators=["\n\n", "\n", ". ", " ", ""],
            )
            docx_doc = Document(page_content=full_text, metadata={"source": docx_path.name})
            docx_chunks = splitter.split_documents([docx_doc])
            
            # Chuẩn hóa metadata
            for idx, chunk in enumerate(docx_chunks):
                chunk.metadata = {
                    "source": docx_path.name,
                    "chunk": idx + 1,
                    "type": "DOCX"
                }
            all_chunks.extend(docx_chunks)
            print(f"  → Đã trích xuất {len(docx_chunks)} chunks từ DOCX.")
            
    # 3. Xử lý CSV
    csv_files = list(data_dir.glob("*.csv"))
    if csv_files:
        print("\n--- Đang xử lý các file CSV ---")
        for csv_path in csv_files:
            print(f"  Đang đọc CSV: {csv_path.name}")
            df = pd.read_csv(str(csv_path))
            
            # Chuyển đổi từng dòng dữ liệu thành văn bản có ngữ cảnh ngữ nghĩa
            for idx, row in df.iterrows():
                row_details = []
                for col in df.columns:
                    val = row[col]
                    if pd.notna(val) and str(val).strip() != "":
                        row_details.append(f"{col} là '{val}'")
                
                # Ví dụ: "Dòng 1 trong file sales.csv: Tên khách hàng là 'Minh', Doanh số là '100 triệu', Tháng là '12'"
                contextual_text = f"Dữ liệu dòng {idx + 1} trong file {csv_path.name}: " + ", ".join(row_details)
                chunk = Document(
                    page_content=contextual_text,
                    metadata={
                        "source": csv_path.name,
                        "row": idx + 1,
                        "type": "CSV"
                    }
                )
                all_chunks.extend([chunk])
            print(f"  → Đã xử lý {len(df)} dòng dữ liệu từ CSV thành chunks.")
            
    if not all_chunks:
        raise FileNotFoundError(
            f"Không tìm thấy file PDF, DOCX, hoặc CSV nào trong thư mục '{data_dir}/'. "
            "Vui lòng copy ít nhất một tài liệu vào data/."
        )
        
    return all_chunks

# ── Step 2: Embed & Lưu ChromaDB ──────────────────────────────
def build_vectordb(chunks: list, chroma_path: Path) -> Chroma:
    # Xóa DB cũ nếu đã tồn tại để xây mới hoàn toàn sạch sẽ
    if chroma_path.exists():
        try:
            shutil.rmtree(chroma_path)
            print("  Đã dọn dẹp ChromaDB cũ thành công.")
        except Exception as e:
            print(f"  Không thể xóa ChromaDB cũ: {e}")

    # Thử khởi động text-embedding-004, nếu lỗi tự động fallback sang gemini-embedding-2
    try:
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/text-embedding-004"
        )
        # Test thử nghiệm 1 chuỗi ngắn để kích hoạt lỗi 404 sớm nếu mô hình không khả dụng
        embeddings.embed_query("test")
        print("  ✓ Sử dụng mô hình nhúng: models/text-embedding-004")
    except Exception as e:
        print(f"  ⚠️ models/text-embedding-004 không khả dụng ({e}). Tự động fallback sang models/gemini-embedding-2.")
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/gemini-embedding-2"
        )

    print(f"  Đang tiến hành nhúng (Embedding) {len(chunks)} chunks dữ liệu...")
    try:
        vectordb = Chroma.from_documents(
            documents         = chunks,
            embedding         = embeddings,
            persist_directory = str(chroma_path),
        )
        # Lưu cache local đề phòng quota lỗi
        save_chunks_cache(chunks)
        return vectordb
    except Exception as e:
        print(f"\n  ⚠️ Lỗi nhúng Vector Database ({e}).")
        print("  🔄 Hệ thống tự động lưu trữ cục bộ để sử dụng BM25 Local Retriever (100% Offline & Miễn phí)!")
        save_chunks_cache(chunks)
        return None

def save_chunks_cache(chunks: list):
    cache_path = Path("data/chunks_cache.json")
    serializable = []
    for c in chunks:
        serializable.append({
            "page_content": c.page_content,
            "metadata": c.metadata
        })
    with open(cache_path, "w", encoding="utf-8") as f:
        json.dump(serializable, f, ensure_ascii=False, indent=2)
    print(f"  ✓ Đã sao lưu {len(chunks)} chunks dữ liệu vào cache local tại: {cache_path}")

# ── Step 3: Thử nghiệm tìm kiếm ngữ nghĩa ──────────────────────
def test_retrieval(vectordb: Chroma):
    print("\n" + "="*20 + " THỬ NGHIỆM TÌM KIẾM NGỮ NGHĨA CHROMA " + "="*20)
    query = input("Nhập câu hỏi thử nghiệm tìm kiếm: ")
    results = vectordb.similarity_search(query, k=3)
    display_results(results)

def test_retrieval_local(chunks: list):
    print("\n" + "="*20 + " THỬ NGHIỆM TÌM KIẾM BM25 CỤC BỘ " + "="*20)
    query = input("Nhập câu hỏi thử nghiệm tìm kiếm: ")
    retriever = BM25Retriever.from_documents(chunks)
    retriever.k = 3
    results = retriever.invoke(query)
    display_results(results)

def display_results(results):
    for i, doc in enumerate(results):
        doc_type = doc.metadata.get("type", "unknown")
        source   = doc.metadata.get("source", "unknown")
        
        if doc_type == "PDF":
            page = doc.metadata.get("page", 1)
            citation = f"File PDF: {source} — trang {page}"
        elif doc_type == "CSV":
            row = doc.metadata.get("row", 1)
            citation = f"File CSV: {source} — dòng {row}"
        elif doc_type == "DOCX":
            chunk_num = doc.metadata.get("chunk", 1)
            citation = f"File Word (DOCX): {source} — chunk {chunk_num}"
        else:
            citation = f"Nguồn: {source}"
            
        print(f"\n[Kết quả {i+1}] {citation}")
        print("-" * 50)
        print(doc.page_content.strip())
        print("-" * 50)

# ── Hàm chạy chính ──────────────────────────────────────────
if __name__ == "__main__":
    print("=== KHỞI CHẠY PIPELINE NẠP VÀ HỌC DỮ LIỆU ĐA ĐỊNH DẠNG ===")
    
    try:
        print("\n[Bước 1/3] Nạp và xử lý tài liệu từ thư mục data/... ")
        chunks = load_all_documents(DATA_DIR)
        # Lọc bỏ các chunk trống hoặc chỉ chứa khoảng trắng để tránh lỗi trong langchain-chroma
        chunks = [c for c in chunks if c.page_content and c.page_content.strip()]
        print(f"✓ Tổng số chunks dữ liệu hợp lệ: {len(chunks)}")
        
        print("\n[Bước 2/3] Embedding và xây dựng Vector Database...")
        vectordb = build_vectordb(chunks, CHROMA_PATH)
        
        if vectordb is not None:
            db_size = len(vectordb.get()['ids'])
            print(f"✓ Đã lưu thành công {db_size} vectors vào ChromaDB tại: {CHROMA_PATH.resolve()}")
            # Test (bỏ qua nếu chạy ngầm từ Streamlit)
            if "--no-test" not in sys.argv:
                test_retrieval(vectordb)
        else:
            print("\n✓ Đã lưu thành công dữ liệu vào cache cục bộ! Sẵn sàng chạy BM25 Local Retriever.")
            # Test retrieval local (bỏ qua nếu chạy ngầm từ Streamlit)
            if "--no-test" not in sys.argv:
                test_retrieval_local(chunks)
        
        print("\n✅ Quá trình Indexing hoàn tất! Dữ liệu đã sẵn sàng cho RAG Chatbot.")
        
    except FileNotFoundError as e:
        print(f"\n❌ Lỗi: {e}")
    except Exception as e:
        import traceback
        print("\n❌ Đã xảy ra lỗi không mong muốn:")
        traceback.print_exc()